"""Analyze DOCX documents to detect page length and content metrics."""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches


class DocxAnalyzer:
    """Analyze DOCX documents for page length, content density, and structure."""

    # Standard US Letter dimensions
    STANDARD_PAGE_HEIGHT_INCHES = 11.0
    STANDARD_PAGE_WIDTH_INCHES = 8.5

    # Margin defaults
    DEFAULT_TOP_MARGIN = 0.5
    DEFAULT_BOTTOM_MARGIN = 0.5
    DEFAULT_LEFT_MARGIN = 0.5
    DEFAULT_RIGHT_MARGIN = 0.5

    # Density thresholds (approx lines per page)
    LINES_PER_PAGE = 50  # Conservative estimate for single-spaced text

    def __init__(self, docx_path: str | Path):
        """Initialize analyzer with a DOCX file.
        
        Args:
            docx_path: Path to the DOCX file
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file isn't a valid DOCX
        """
        self.docx_path = Path(docx_path)
        if not self.docx_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")

        try:
            self.doc = Document(str(self.docx_path))
        except Exception as e:
            raise ValueError(f"Failed to open DOCX file: {e}")

    def estimate_page_count(self) -> int:
        """Estimate the number of pages in the document.
        
        Returns:
            Estimated page count (1, 2, or 3)
        """
        line_count = self._count_content_lines()
        pages = max(1, (line_count + self.LINES_PER_PAGE - 1) // self.LINES_PER_PAGE)
        return min(3, pages)  # Cap at 3 pages

    def _count_content_lines(self) -> int:
        """Count approximate lines of content."""
        total_lines = 0

        for paragraph in self.doc.paragraphs:
            if not paragraph.text.strip():
                continue

            # Estimate lines based on character count and average line width
            # Average chars per line varies, but ~80-100 is reasonable for US Letter
            chars = len(paragraph.text)
            lines = max(1, (chars + 90) // 90)  # Rough estimate
            total_lines += lines

            # Add extra lines for spacing after paragraphs
            if paragraph.paragraph_format.space_after:
                space_points = paragraph.paragraph_format.space_after.pt
                total_lines += space_points / 12  # Rough conversion to line count

        return int(total_lines)

    def get_word_count(self) -> int:
        """Get total word count in document."""
        total_words = 0
        for paragraph in self.doc.paragraphs:
            total_words += len(paragraph.text.split())
        return total_words

    def get_section_analysis(self) -> dict[str, dict]:
        """Analyze content by sections (experience, projects, skills, etc).
        
        Returns:
            Dictionary with section names and their metrics
        """
        sections = {}
        current_section = "other"
        section_content = {"word_count": 0, "item_count": 0, "lines": 0}

        # Common section headers (case-insensitive)
        section_headers = {
            "experience": ["experience", "employment", "work history", "professional"],
            "projects": ["projects", "portfolio", "side projects"],
            "skills": ["skills", "technical skills", "competencies"],
            "education": ["education", "degree", "university"],
            "certifications": ["certifications", "certificates", "certs"],
            "awards": ["awards", "honors", "recognition"],
        }

        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            # Check if this is a section header
            lower_text = text.lower()
            is_header = False

            for section_name, headers in section_headers.items():
                if any(h in lower_text for h in headers):
                    if current_section != "other":
                        sections[current_section] = section_content
                    current_section = section_name
                    section_content = {"word_count": 0, "item_count": 0, "lines": 0}
                    is_header = True
                    break

            if not is_header:
                word_count = len(text.split())
                section_content["word_count"] += word_count
                section_content["lines"] += max(1, (len(text) + 90) // 90)

                # Count items (bullets, numbered lists)
                if paragraph.style.name.startswith("List"):
                    section_content["item_count"] += 1

        # Save last section
        if current_section != "other" or section_content["word_count"] > 0:
            sections[current_section] = section_content

        return sections

    def predict_page_impact(self, word_reduction: int) -> int:
        """Predict page count after removing words.
        
        Args:
            word_reduction: Number of words to remove
            
        Returns:
            Predicted page count after reduction
        """
        current_words = self.get_word_count()
        new_words = max(100, current_words - word_reduction)  # Minimum 100 words

        # Rough estimate: ~250-300 words per page for a condensed CV
        words_per_page = 275
        estimated_pages = max(1, (new_words + words_per_page - 1) // words_per_page)

        return min(3, estimated_pages)

    def get_content_density(self) -> float:
        """Get content density (words per page estimate).
        
        Returns:
            Estimated words per page
        """
        words = self.get_word_count()
        pages = self.estimate_page_count()
        return words / max(1, pages)

    def analyze_spacing(self) -> dict[str, float]:
        """Analyze paragraph spacing (can be optimized for space savings).
        
        Returns:
            Dictionary with spacing metrics that could be optimized
        """
        total_before = 0.0
        total_after = 0.0
        total_line_spacing = 0.0
        paragraph_count = 0

        for paragraph in self.doc.paragraphs:
            if paragraph.text.strip():
                pf = paragraph.paragraph_format
                total_before += (pf.space_before.pt if pf.space_before else 0)
                total_after += (pf.space_after.pt if pf.space_after else 0)
                total_line_spacing += (pf.line_spacing if pf.line_spacing else 1.0)
                paragraph_count += 1

        if paragraph_count == 0:
            return {"avg_space_before": 0, "avg_space_after": 0, "avg_line_spacing": 1.0}

        return {
            "avg_space_before": total_before / paragraph_count,
            "avg_space_after": total_after / paragraph_count,
            "avg_line_spacing": total_line_spacing / paragraph_count,
            "paragraph_count": paragraph_count,
        }
