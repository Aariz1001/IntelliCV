from __future__ import annotations

from typing import Iterable, Optional, Any
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from docx.enum.section import WD_SECTION, WD_ORIENT
from .cv_config import DocxFormat


def _set_run_font(run, name: str = "Lora", size: int = 10, bold: bool = False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold


def _set_paragraph_spacing(paragraph, space_before: int = 0, space_after: int = 0, line_spacing: float = 1.0):
    """Set paragraph spacing in points."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = line_spacing


def _add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_heading(doc: Document, text: str, docx_format: Optional[DocxFormat] = None):
    if docx_format is None:
        docx_format = DocxFormat()
    
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text.upper())
    _set_run_font(run, name=docx_format.font_family, size=docx_format.font_sizes.section_heading, bold=docx_format.formatting.section_headers_bold)
    _set_paragraph_spacing(paragraph, space_before=docx_format.spacing.section_before, space_after=docx_format.spacing.section_after)
    
    if docx_format.formatting.section_headers_underlined:
        _add_horizontal_line(paragraph)


def _add_bullets(doc: Document, bullets: Iterable[str], font_size: int = 9, docx_format: Optional[DocxFormat] = None):
    if docx_format is None:
        docx_format = DocxFormat()
    
    for bullet in bullets:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(bullet)
        _set_run_font(run, name=docx_format.font_family, size=font_size)
        _set_paragraph_spacing(paragraph, space_before=docx_format.spacing.bullet_before, space_after=docx_format.spacing.bullet_after)


def _setup_document(doc: Document, docx_format: Optional[DocxFormat] = None):
    """Set up page with specified format (or A4 with 0.25" margins by default)."""
    if docx_format is None:
        docx_format = DocxFormat()  # Use defaults
    
    for section in doc.sections:
        # Set page size
        section.page_height = Inches(docx_format.page_height_inches)
        section.page_width = Inches(docx_format.page_width_inches)
        
        # Set margins
        section.top_margin = Inches(docx_format.margin_inches)
        section.bottom_margin = Inches(docx_format.margin_inches)
        section.left_margin = Inches(docx_format.margin_inches)
        section.right_margin = Inches(docx_format.margin_inches)
        
        # Remove header/footer distances
        section.header_distance = Inches(0)
        section.footer_distance = Inches(0)
        
        # Clear any headers/footers
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True
        for para in section.header.paragraphs:
            para.clear()
        for para in section.footer.paragraphs:
            para.clear()


def build_docx(cv: dict, output_path: str, docx_format: Optional[DocxFormat] = None) -> None:
    if docx_format is None:
        docx_format = DocxFormat()
    
    doc = Document()
    _setup_document(doc, docx_format)

    # Name - centered, larger
    name_para = doc.add_paragraph()
    name_para.alignment = 1  # Center
    name_run = name_para.add_run(cv.get("name", ""))
    _set_run_font(name_run, name=docx_format.font_family, size=docx_format.font_sizes.name, bold=docx_format.formatting.name_bold)
    _set_paragraph_spacing(name_para, space_before=docx_format.spacing.name_before, space_after=docx_format.spacing.name_after)

    # Title
    title = cv.get("title")
    if title:
        title_para = doc.add_paragraph()
        title_para.alignment = 1  # Center
        title_run = title_para.add_run(title)
        _set_run_font(title_run, name=docx_format.font_family, size=docx_format.font_sizes.title)
        _set_paragraph_spacing(title_para, space_before=0, space_after=0)

    # Contact info - all on one line
    contact = cv.get("contact", {})
    contact_parts = [
        item for item in [
            contact.get("email"),
            contact.get("phone"),
            contact.get("location"),
        ] if item
    ]
    links = contact.get("links", [])
    for link in links:
        contact_parts.append(f"{link.get('label', '')}: {link.get('url', '')}")
    
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = 1  # Center
        contact_run = contact_para.add_run(" | ".join(contact_parts))
        _set_run_font(contact_run, name=docx_format.font_family, size=docx_format.font_sizes.contact_info)
        _set_paragraph_spacing(contact_para, space_before=0, space_after=4)

    # Summary
    summary = cv.get("summary", [])
    if summary:
        _add_heading(doc, "Profile", docx_format)
        # Join summary into a single paragraph for compactness
        summary_para = doc.add_paragraph()
        summary_text = " ".join(summary)
        summary_run = summary_para.add_run(summary_text)
        _set_run_font(summary_run, name=docx_format.font_family, size=docx_format.font_sizes.bullet)
        _set_paragraph_spacing(summary_para, space_before=0, space_after=2)

    experience = cv.get("experience", [])
    if experience:
        _add_heading(doc, "Experience", docx_format)
        for item in experience:
            header = f"{item.get('role', '')} — {item.get('company', '')}"
            dates = item.get("dates")
            location = item.get("location", "")
            if location and dates:
                header = f"{header} | {location} | {dates}"
            elif dates:
                header = f"{header} | {dates}"
            elif location:
                header = f"{header} | {location}"
            header_para = doc.add_paragraph()
            header_run = header_para.add_run(header)
            _set_run_font(header_run, name=docx_format.font_family, size=docx_format.font_sizes.role_header, bold=docx_format.formatting.role_headers_bold)
            _set_paragraph_spacing(header_para, space_before=docx_format.spacing.role_before, space_after=docx_format.spacing.role_after)
            _add_bullets(doc, item.get("bullets", []), font_size=docx_format.font_sizes.bullet, docx_format=docx_format)

    projects = cv.get("projects", [])
    if projects:
        _add_heading(doc, "Projects", docx_format)
        for item in projects:
            project_para = doc.add_paragraph()
            # Project name in bold
            name_run = project_para.add_run(item.get("name", ""))
            _set_run_font(name_run, name=docx_format.font_family, size=docx_format.font_sizes.bullet, bold=True)
            # Add description as paragraph text (not bullets)
            bullets = item.get("bullets", [])
            if bullets:
                desc_text = " ".join(bullets)  # Combine all bullets into one paragraph
                desc_run = project_para.add_run(f": {desc_text}")
                _set_run_font(desc_run, name=docx_format.font_family, size=docx_format.font_sizes.bullet)
            _set_paragraph_spacing(project_para, space_before=2, space_after=2)

    education = cv.get("education", [])
    if education:
        _add_heading(doc, "Education", docx_format)
        for item in education:
            header = f"{item.get('degree', '')} — {item.get('school', '')}"
            location = item.get("location", "")
            dates = item.get("dates", "")
            if location and dates:
                header = f"{header} | {location} | {dates}"
            elif dates:
                header = f"{header} | {dates}"
            elif location:
                header = f"{header} | {location}"
            header_para = doc.add_paragraph()
            header_run = header_para.add_run(header)
            _set_run_font(header_run, name=docx_format.font_family, size=docx_format.font_sizes.role_header, bold=docx_format.formatting.role_headers_bold)
            _set_paragraph_spacing(header_para, space_before=2, space_after=1)
            details = item.get("details", [])
            if details:
                _add_bullets(doc, details, font_size=docx_format.font_sizes.bullet, docx_format=docx_format)

    skills = cv.get("skills", {}).get("groups", [])
    if skills:
        _add_heading(doc, "Skills", docx_format)
        # Combine all skills into compact format
        for group in skills:
            group_para = doc.add_paragraph()
            group_name = group.get("name", "")
            items = ", ".join(group.get("items", []))
            group_run = group_para.add_run(f"{group_name}: ")
            _set_run_font(group_run, name=docx_format.font_family, size=docx_format.font_sizes.skills_label, bold=True)
            items_run = group_para.add_run(items)
            _set_run_font(items_run, name=docx_format.font_family, size=docx_format.font_sizes.skills_label)
            _set_paragraph_spacing(group_para, space_before=0, space_after=1)

    certifications = cv.get("certifications", [])
    if certifications:
        _add_heading(doc, "Certifications", docx_format)
        cert_para = doc.add_paragraph()
        cert_run = cert_para.add_run(" | ".join(certifications))
        _set_run_font(cert_run, size=9)
        _set_paragraph_spacing(cert_para, space_before=0, space_after=1)

    awards = cv.get("awards", [])
    if awards:
        _add_heading(doc, "Awards")
        awards_para = doc.add_paragraph()
        awards_run = awards_para.add_run(" | ".join(awards))
        _set_run_font(awards_run, size=9)
        _set_paragraph_spacing(awards_para, space_before=0, space_after=1)

    doc.save(output_path)
