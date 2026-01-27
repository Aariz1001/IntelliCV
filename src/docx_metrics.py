"""Extract and store metrics from DOCX documents."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Optional, Any

from docx import Document
from docx.oxml.ns import qn


class DocxMetricsExtractor:
    """Extract formatting and content metrics from DOCX files."""
    
    def __init__(self, docx_path: str | Path):
        """Initialize with a DOCX file."""
        self.docx_path = Path(docx_path)
        self.doc = Document(str(self.docx_path))
    
    def extract_all_metrics(self) -> dict[str, Any]:
        """Extract all available metrics from the document."""
        return {
            "word_count": self._count_words(),
            "page_count": self._estimate_pages(),
            "source_format": "docx",
            "extracted_fonts": self._extract_fonts(),
            "margins": self._extract_margins(),
            "extraction_timestamp": datetime.now().isoformat(),
        }
    
    def _count_words(self) -> int:
        """Count total words in document."""
        total = 0
        for paragraph in self.doc.paragraphs:
            total += len(paragraph.text.split())
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    total += len(cell.text.split())
        return total
    
    def _estimate_pages(self) -> int:
        """Calculate page count more accurately based on document content and formatting."""
        try:
            section = self.doc.sections[0]
            # Available height in inches
            page_height = section.page_height.inches
            top_margin = section.top_margin.inches
            bottom_margin = section.bottom_margin.inches
            available_height = page_height - top_margin - bottom_margin
            
            # Calculate total content height
            # Default line height is approximately 1/6 inch (12pt)
            # With spacing between paragraphs
            total_height = 0.0
            
            for paragraph in self.doc.paragraphs:
                if paragraph.text.strip():
                    # Get paragraph formatting
                    para_format = paragraph.paragraph_format
                    
                    # Space before and after paragraph (in inches)
                    space_before = para_format.space_before.inches if para_format.space_before else 0
                    space_after = para_format.space_after.inches if para_format.space_after else 0
                    
                    # Get line spacing (default 1.0 = single spacing)
                    line_spacing = para_format.line_spacing if para_format.line_spacing else 1.0
                    if line_spacing < 0.1:  # If it's a point value
                        line_spacing = 1.0
                    
                    # Get font size (default 12pt = 1/6 inch)
                    font_size_pt = 12
                    for run in paragraph.runs:
                        if run.font.size:
                            font_size_pt = int(run.font.size.pt)
                            break
                    
                    # One line height in inches (12pt ≈ 0.167 inches)
                    line_height = (font_size_pt / 72.0) * line_spacing
                    
                    # Estimate lines needed for this paragraph
                    # Approximate: ~80 characters per line at 1" width
                    available_width = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
                    chars_per_line = int(available_width * 12)  # Rough estimate: 12 chars per inch
                    lines_needed = max(1, (len(paragraph.text) + chars_per_line - 1) // chars_per_line)
                    
                    # Total height for this paragraph
                    para_height = space_before + (lines_needed * line_height) + space_after
                    total_height += para_height
            
            # Account for tables
            for table in self.doc.tables:
                # Rough estimate: each table row is ~0.3 inches
                table_height = len(table.rows) * 0.3
                total_height += table_height
            
            # Calculate pages (minimum 1)
            pages = max(1, int((total_height + available_height - 1) / available_height))
            return pages
            
        except Exception:
            # Fallback to word-based estimation if calculation fails
            word_count = self._count_words()
            pages = max(1, (word_count + 249) // 250)
            return pages
    
    def _extract_fonts(self) -> dict[str, Any]:
        """Extract font information from document."""
        fonts = {}
        font_sizes = {}
        
        # Sample fonts from document
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                if run.font.name:
                    fonts[run.font.name] = fonts.get(run.font.name, 0) + 1
                if run.font.size:
                    size_pt = int(run.font.size.pt)
                    font_sizes[str(size_pt)] = font_sizes.get(str(size_pt), 0) + 1
        
        # Get most common font
        primary_font = max(fonts.items(), key=lambda x: x[1])[0] if fonts else "Calibri"
        
        # Sort detected sizes by frequency (descending)
        sorted_sizes = sorted(
            font_sizes.items(),
            key=lambda x: x[1],
            reverse=True
        )
        
        return {
            "family": primary_font,
            "detected_sizes": {
                size: count for size, count in sorted_sizes
            }
        }
    
    def _extract_margins(self) -> dict[str, float]:
        """Extract page margins."""
        try:
            section = self.doc.sections[0]
            return {
                "top": section.top_margin.inches,
                "bottom": section.bottom_margin.inches,
                "left": section.left_margin.inches,
                "right": section.right_margin.inches,
            }
        except Exception:
            # Return defaults if can't extract
            return {
                "top": 0.5,
                "bottom": 0.5,
                "left": 0.5,
                "right": 0.5,
            }
